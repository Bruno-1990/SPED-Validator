"""Conversor de documentos SPED (PDF, DOCX, TXT) para Markdown estruturado."""

from __future__ import annotations

import re
from pathlib import Path
from typing import Any

import pdfplumber
from tqdm import tqdm

# Threshold para detectar headings (fonte maior que o corpo)
_HEADING_SIZE_DELTA = 1.5

# Extensoes suportadas
SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt"}


# ──────────────────────────────────────────────
# API publica
# ──────────────────────────────────────────────

def convert_file_to_markdown(file_path: str | Path) -> str:
    """Converte qualquer arquivo suportado (PDF, DOCX, TXT) para Markdown."""
    file_path = Path(file_path)
    ext = file_path.suffix.lower()

    if ext == ".pdf":
        return _convert_pdf(file_path)
    elif ext == ".docx":
        return _convert_docx(file_path)
    elif ext == ".txt" or ext == "":
        return _convert_txt(file_path)
    else:
        raise ValueError(f"Formato nao suportado: {ext}")


def convert_all_docs(
    input_dir: str | Path,
    output_dir: str | Path,
    skip_existing: bool = True,
) -> list[Path]:
    """Converte todos os arquivos suportados de um diretorio para Markdown.

    Retorna lista de arquivos .md criados.
    """
    input_dir = Path(input_dir)
    output_dir = Path(output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)

    files = _find_supported_files(input_dir)
    created: list[Path] = []

    for file_path in tqdm(files, desc=f"Convertendo {input_dir.name}"):
        md_path = output_dir / f"{file_path.stem}.md"

        if skip_existing and md_path.exists():
            continue

        try:
            markdown = convert_file_to_markdown(file_path)
            if markdown.strip():
                md_path.write_text(markdown, encoding="utf-8")
                created.append(md_path)
        except Exception as e:
            print(f"  ERRO ao converter {file_path.name}: {e}")

    return created


def _find_supported_files(directory: Path) -> list[Path]:
    """Encontra todos os arquivos suportados no diretorio (inclui sem extensao)."""
    files: list[Path] = []
    for ext in SUPPORTED_EXTENSIONS:
        files.extend(directory.glob(f"*{ext}"))
        files.extend(directory.glob(f"*{ext.upper()}"))
    # Incluir arquivos sem extensao (tratados como txt)
    for f in directory.iterdir():
        if f.is_file() and f.suffix == "" and f not in files:
            files.append(f)
    return sorted(set(files))


# ──────────────────────────────────────────────
# Conversor PDF
# ──────────────────────────────────────────────

def _convert_pdf(pdf_path: Path) -> str:
    """Converte um PDF em Markdown estruturado."""
    sections: list[str] = []
    sections.append(f"# {pdf_path.stem}\n")

    with pdfplumber.open(pdf_path) as pdf:
        body_size = _estimate_body_font_size(pdf)

        for page_num, page in enumerate(pdf.pages, start=1):
            page_md = _process_page(page, body_size, page_num)
            if page_md.strip():
                sections.append(page_md)

    return "\n\n".join(sections)


# ──────────────────────────────────────────────
# Conversor DOCX
# ──────────────────────────────────────────────

def _convert_docx(docx_path: Path) -> str:
    """Converte um DOCX em Markdown."""
    try:
        from docx import Document
    except ImportError:
        # Fallback: extrair texto bruto via zipfile
        return _convert_docx_fallback(docx_path)

    doc = Document(str(docx_path))
    sections: list[str] = []
    sections.append(f"# {docx_path.stem}\n")

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        style_name = (para.style.name or "").lower()

        if "heading 1" in style_name:
            sections.append(f"\n## {text}\n")
        elif "heading 2" in style_name:
            sections.append(f"\n### {text}\n")
        elif "heading 3" in style_name:
            sections.append(f"\n#### {text}\n")
        elif "title" in style_name:
            sections.append(f"\n## {text}\n")
        else:
            sections.append(text)

    # Extrair tabelas
    for table in doc.tables:
        md_table = _docx_table_to_markdown(table)
        if md_table:
            sections.append(f"\n{md_table}\n")

    return "\n\n".join(sections)


def _convert_docx_fallback(docx_path: Path) -> str:
    """Extrai texto de DOCX sem python-docx (via zipfile)."""
    import xml.etree.ElementTree as ET  # nosec B405 - DOCX local, não XML externo
    import zipfile

    sections: list[str] = [f"# {docx_path.stem}\n"]

    with zipfile.ZipFile(docx_path, "r") as z:
        if "word/document.xml" not in z.namelist():
            return ""
        with z.open("word/document.xml") as f:
            tree = ET.parse(f)  # nosec B314 # noqa: S314 - DOCX local

    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    for para in tree.iter(f"{{{ns['w']}}}p"):
        texts = [t.text for t in para.iter(f"{{{ns['w']}}}t") if t.text]
        line = "".join(texts).strip()
        if line:
            sections.append(line)

    return "\n\n".join(sections)


def _docx_table_to_markdown(table: object) -> str:
    """Converte tabela do python-docx para Markdown."""
    rows_data: list[list[str]] = []
    for row in table.rows:  # type: ignore[attr-defined]
        cells = [cell.text.strip().replace("|", "\\|") for cell in row.cells]
        rows_data.append(cells)

    if len(rows_data) < 2:
        return ""

    max_cols = max(len(r) for r in rows_data)
    for row in rows_data:
        while len(row) < max_cols:
            row.append("")

    header = rows_data[0]
    lines = [
        "| " + " | ".join(header) + " |",
        "| " + " | ".join("---" for _ in header) + " |",
    ]
    for row in rows_data[1:]:
        lines.append("| " + " | ".join(row) + " |")

    return "\n".join(lines)


# ──────────────────────────────────────────────
# Conversor TXT
# ──────────────────────────────────────────────

def _convert_txt(txt_path: Path) -> str:
    """Converte um TXT em Markdown."""
    # Tentar varios encodings
    for enc in ("utf-8", "latin-1", "cp1252"):
        try:
            content = txt_path.read_text(encoding=enc)
            break
        except (UnicodeDecodeError, LookupError):
            continue
    else:
        content = txt_path.read_bytes().decode("latin-1", errors="replace")

    sections: list[str] = [f"# {txt_path.stem}\n"]
    sections.append(content)
    return "\n\n".join(sections)


# ──────────────────────────────────────────────
# Funcoes auxiliares PDF
# ──────────────────────────────────────────────

def _estimate_body_font_size(pdf: pdfplumber.PDF) -> float:
    """Estima o tamanho de fonte do corpo analisando as primeiras paginas."""
    sizes: list[float] = []
    for page in pdf.pages[:5]:
        if page.chars:
            sizes.extend(c.get("size", 0) for c in page.chars if c.get("size"))

    if not sizes:
        return 10.0

    from collections import Counter
    rounded = [round(s, 1) for s in sizes]
    most_common = Counter(rounded).most_common(1)
    return most_common[0][0] if most_common else 10.0


def _process_page(page: Any, body_size: float, page_num: int) -> str:
    """Processa uma pagina do PDF extraindo texto e tabelas."""
    parts: list[str] = []

    tables = page.find_tables()
    table_bboxes = [t.bbox for t in tables]

    text_outside = _extract_text_outside_tables(page, table_bboxes)

    if text_outside:
        lines_with_type = _classify_lines(page, body_size, table_bboxes)
        for line_text, line_type in lines_with_type:
            line_text = line_text.strip()
            if not line_text:
                continue
            if line_type == "heading":
                level = _detect_heading_level(line_text)
                parts.append(f"\n{'#' * level} {line_text}\n")
            else:
                parts.append(line_text)

    for table in tables:
        extracted = table.extract()
        if extracted:
            md_table = _table_to_markdown(extracted)
            if md_table:
                parts.append(f"\n{md_table}\n")

    return "\n".join(parts)


def _extract_text_outside_tables(
    page: Any,
    table_bboxes: list[tuple],
) -> str:
    """Extrai texto da pagina excluindo areas de tabelas."""
    if not table_bboxes:
        return page.extract_text() or ""

    filtered = page
    for bbox in table_bboxes:
        filtered = filtered.outside_bbox(bbox)

    return filtered.extract_text() or ""


def _classify_lines(
    page: Any,
    body_size: float,
    table_bboxes: list[tuple],
) -> list[tuple[str, str]]:
    """Classifica cada linha do texto como 'heading' ou 'body'."""
    chars = page.chars
    if not chars:
        text = page.extract_text() or ""
        return [(line, "body") for line in text.split("\n")]

    filtered_chars = [c for c in chars if not _char_in_bbox(c, table_bboxes)]

    if not filtered_chars:
        return []

    lines = _group_chars_into_lines(filtered_chars)
    result: list[tuple[str, str]] = []

    for line_chars in lines:
        text = "".join(c.get("text", "") for c in line_chars)
        avg_size = sum(c.get("size", body_size) for c in line_chars) / len(line_chars)

        if avg_size > body_size + _HEADING_SIZE_DELTA:
            result.append((text, "heading"))
        else:
            result.append((text, "body"))

    return result


def _group_chars_into_lines(chars: list[dict]) -> list[list[dict]]:
    """Agrupa caracteres em linhas baseado na posicao vertical."""
    if not chars:
        return []

    sorted_chars = sorted(chars, key=lambda c: (round(c.get("top", 0), 1), c.get("x0", 0)))

    lines: list[list[dict]] = []
    current_line: list[dict] = [sorted_chars[0]]
    current_top = round(sorted_chars[0].get("top", 0), 1)

    for char in sorted_chars[1:]:
        char_top = round(char.get("top", 0), 1)
        if abs(char_top - current_top) < 2:
            current_line.append(char)
        else:
            lines.append(current_line)
            current_line = [char]
            current_top = char_top

    if current_line:
        lines.append(current_line)

    return lines


def _char_in_bbox(char: dict, bboxes: list[tuple]) -> bool:
    """Verifica se um caractere esta dentro de alguma bounding box."""
    cx = char.get("x0", 0)
    cy = char.get("top", 0)
    return any(x0 <= cx <= x1 and y0 <= cy <= y1 for x0, y0, x1, y1 in bboxes)


def _detect_heading_level(text: str) -> int:
    """Detecta o nivel do heading baseado no conteudo."""
    text_upper = text.upper().strip()

    if re.match(r"^BLOCO\s+[A-Z0-9]", text_upper):
        return 2
    if re.match(r"^REGISTRO\s+[A-Z0-9]{4}", text_upper):
        return 3
    if any(kw in text_upper for kw in ("CAMPOS DO REGISTRO", "CAMPO", "TABELA")):
        return 4
    if any(kw in text_upper for kw in ("OBSERVA", "NOTA", "REGRA")):
        return 4

    return 3


def _table_to_markdown(table_data: list[list[str | None]]) -> str:
    """Converte dados de tabela extraidos em formato Markdown."""
    if not table_data or len(table_data) < 2:
        return ""

    cleaned: list[list[str]] = [
        [_clean_cell(cell) for cell in row]
        for row in table_data
    ]

    max_cols = max(len(row) for row in cleaned)
    for row in cleaned:
        while len(row) < max_cols:
            row.append("")

    header = cleaned[0]
    lines = [
        "| " + " | ".join(header) + " |",
        "| " + " | ".join("---" for _ in header) + " |",
    ]

    for row in cleaned[1:]:
        lines.append("| " + " | ".join(row) + " |")

    return "\n".join(lines)


def _clean_cell(cell: str | None) -> str:
    """Limpa o conteudo de uma celula de tabela."""
    if cell is None:
        return ""
    text = re.sub(r"\s+", " ", str(cell)).strip()
    text = text.replace("|", "\\|")
    return text
